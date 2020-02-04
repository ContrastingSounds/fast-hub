import datetime

from docx import Document
from docx.shared import Inches
from fastapi import BackgroundTasks

from main import app
from core import get_sdk_all_access, send_email, get_temp_file_name, get_output_file_name


slug = 'folder_to_word'

def write_docx_from_folder(folder_id: int):
    sdk = get_sdk_all_access()
    folder = sdk.folder(folder_id)
    looks = folder.looks

    timestamp = datetime.datetime.now()

    document = Document()
    document.sections[0].header.paragraphs[0].text = folder.name
    document.sections[0].footer.paragraphs[0].text = f'Created {timestamp}'

    titles = []
    pages = len(looks) - 1

    for idx, look in enumerate(looks):
        image = sdk.run_look(look.id, 'png')
        image_file = get_temp_file_name(
            slug, 
            '.'.join(['look', str(look.id), 'png'])
        )
        with open(image_file, 'wb') as file:
            file.write(image)
        document.add_heading(look.title, 0)
        document.add_paragraph(look.description)
        document.add_picture(image_file, width=Inches(5))
        if idx < pages:
            document.add_page_break()
    
    word_file_name = get_output_file_name(
        slug, 
        '.'.join([folder.name, 'docx']),
        timestamp=True
    )
    document.save(word_file_name)


@app.get('/extensions/%s/folder/{folder_id}' % slug, status_code=202)
def endpoint(folder_id: int, background_tasks: BackgroundTasks):
    """Creates Word document from all looks in a given folder"""
    background_tasks.add_task(write_docx_from_folder, folder_id)

    return {'message': 'Generating DOCS in background'}


@app.get(f'/extensions/{slug}/folders', status_code=200)
def folders():
    """Folders endpoint for Folder to Word action: returns list of all folders. Does not include root
    folders i.e. Shared, Users and LookML."""
    sdk = get_sdk_all_access()
    all_folders = sdk.all_folders()

    folders = []
    for folder in all_folders:
        if folder.name not in ['Shared', 'Users', 'lookml']:
            folders.append({
                folder.id: folder.name
            })

    return {'folders': folders}
