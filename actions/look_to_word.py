import base64
import csv
import datetime

from docx import Document
from docx.shared import Inches

from main import app
from core import action_hub, get_output_file_name, get_temp_file_name, get_sdk_for_schedule, get_sdk_all_access, send_email
from api_types import ActionDefinition, ActionList, ActionFormField, ActionRequest, ActionForm

slug = 'look_to_word'

icon_data_uri = 'data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAOEAAADhCAMAAAAJbSJIAAAAh1BMVEUrV5r///8XTZUdUJfy9fkoVZkkU5ggUZcARpJffK+On8ISS5Sdrcvv8fbFz+CZqslFZ6J/k7o6YqH5+v2TpcbU2+isudIISJPN1eQAQ5Fuh7To7PPg5e5pg7KCl72jsc25xNlSc6m/ydxthrQ0XZ2xvdTa4exOb6dgfa9/lLsAP5BYdqt2jrkmH1BUAAAKgElEQVR4nO2d6YKiOhCFkb3FbURF27V73Hra93++y+JSVSlQGyRJX87PGQbymXCqUkkYo/XbZchuwMvVEOqvhlB/SSQMF5vR0Xn5Y+QQzqfr0zCIAssLXv6sugnD8a79bfiB6TlGLMd8+RPrIwwXn/2/ziwyLde46rcQdqeTr4Prx4PSINKf0N5v2kcvGZQuhdOfcN7rnIZWFOSwaU1oj3erbyvuONdxCuj0JFx8jo5GELvJXTb9CJMQdzBjN3EfY9OJMIzd5Dt+4WI3eQZOD8LYTf4OktzkeTblCe3pZHXwr7nJT6Um4T5OmLexmxSGAU0J5731aeCXGJQKEyYJ84cTlR2UShImCfNy6CchrlI2JQi70/XXwEpmca+Ak0sY7nftt63/oxCnPGGSMMdOWZwwa0pojyerDydNmOuAq5UwdpPR0pq9yE1kE8YJ88BKEuaa2eojfL2byCaUK0QYfh/fKtWyvVCMMDbwamX926lF+AKjm40rv+Vzejmhu6z8ls/p5YTS1RDqr4ZQfzWE+qsh1F8Nof5qCPVXQ6i/GkL91RDqL0wYmVY5KfgLYcJ+p5zWR/UQK6559y3ZQIIqJhw1hPWrIWwIG0L5+r8RdkvKbitOGP7zS0o9wP9d5q07oWOIALoTOq7rep5pBkEURaZnbMUr9CZ0jMPx9DVa7za98WLeDVthIFyiMmE86BwvlhV3UcDuqDLXJFz1dCKMZoEx+FieVu3+ZPM5PXGIwefdgCybMOkml+2eaI/b/in0TnLRghB+CASyCK3EGeKxZxrD+E1iHutYIW77wmduI5wPFS+SRGj1p5kzZDqIvegcaOO5n2FLrtlH4jWY8BU0nMwdatgf4VykYa0o4bf4M3jv5JqJKVyDCX2zpB7dQuqdUMPWYsOCDSVkcl6rT65h7Ahn3nZJhY9m3s4ANUz0eNFDWhvxIsFKh3dymvJ6ePYUIR9ZiK/PTLg3cxH12y7jRtIIx/CfiRYnGk2rNRPuEhC/ZcaCNEJzgv7dgBJabfHmguM6wwceL4uQWM2SmqloNK3Wit7cPdLfgIsokgjJzy/8O18wmlZrR8egNcIXiGm3RELDtAsb74bizYVwHuCo2hqLXiSREFsNbZvzwdxc8CN8Dzbey+xDZDU2sXnOaESr8W389+9cwiEp4gtWY+Lu4YwmthrsR45F/n7LpdWSsjYhqyFJJ2c0wih03/Bfz7nZh6zMO5aPvIREgogDpFZDhzI7g5Q4A8Y2gbvH+WYJSVZDJihcam7IJMRWM0W/P280NPWJpvhvmVmmIZPQ+wNbh98hYDQL6Dl4Hhl1EWBosq2XR4izmhD9/mDq9Am7k1gN7kI23kslxFnNACKCjKYPDRNB0KyAmUZnz5FFSKwGRmvY9ndYbLLhQPRIneNdLIWkN5NHiK2mA4zQAxm1M4NvG5zD02pwzmMkEmKrgcEMGE03CqBjnkA/BdhK+XhPCZPPLpXTE4TYamA0B0YzDVBXwXfNnyNCPt5TwuGgpNjMME8zmNWEt2gO8821iRJYGDZJ3pOXEsvc9YWt5vbnMN98d1ECC+YgLqnkDHJyYpk7FbDVHK8thFP32FpQuelmNWRyYud5iExC3MbbKANV0KTPUFffggqx0px4L5cQT6B2VxOJbhaSvHcowb5ZDakG58R7yYQoq5leO8HD7UYrGLeKKCmKC9U6FQjx+Ote2u6CqVOyDoGKw/PLVY6Ha1W5D5ZKiK3m4hTQaJLo46BFwst7SMqRefFeMiHOai7TO/CCZf2KxuPb+SoXL6wxyzYqEOJ+uEz+QK6SvZuoLHV5AllYE+rhOYS9kpqe8t53Xiiruboh/SNU2r50FrHSvHhPs7ZZUFLPAWKrOdskXIzIoh9anrgs5+OFNZuf34uEde/FQFZzbjvJaFJBmDMLLtVN8+K9bEKc1WSEJKNJhGYRmSERK+3kxXvZhC7KarIeIxlNoqAHrvpK3wRviQiX+aVoyTuGUFadpiV46pRdhXxzIrpPqxXlN1wyIbKaNBAgozkbl/cXXDUWIwi3xH+VZEJkNWkggP11m0+Dq7opDbbS/HgvndD7or0D3jn72u4AVqMc4U/oohSSZEJkiWmOFtGMJhEqOmWzCAhYEO+lE+J9NQ7eqXarcKPZblIpwQtrdsEgFWptJfU8IbSaDwcZzS0FRPvXkqIaXrqZPk748V1Ob8NnCZHVxOkzNJrbShPag5gYJ15Y6xTl+7JPWKKsJo5/IKMB9UW0Yh+6DinTHYvGjuxTQSiriUcbSNBgbQlZTZy3+chK89NuBQiR1cx9B5R54VIaspokb4MPLYr3KhDC8WZ5wCPhciia0U9MF62CC9uN1CJEVjMwQboJl7RR3BxHeGHtq3BWKp0Q1WrefTB1Qm8XzGFsH5ewhI2NahGi3un8442GDma0R4Hup1KOEJaFN2A2hdfsUa8dZ3BKXBjv5cdD3Dt7EDvwvgsUN0doi/SdZ2LCUb+cOm/PJ7awd7pgqoE3paM1jg0qYTCnFHIJpXz5A1nNLRkLya5uWHjco2rwnU/yy55bkKzmhkFXy+BgtmFVozjeq0BI6oIX0d2wyGqg0ezyy2ypFCCkO30z0eo53T1zFXOiCEkBQnIw4SwaxnHhEYg5JoOkACFegbo0hQY54bTeWXfivRKEDtc74rK8v2cu44/J4NsrQGjZYsPFbff81u/7WZQChKzViP4hHkhMdSfeq0HIWY24uYo765WVNIpV8V79H/1CqCx8bgkTxn2OkDkWW0TYKv01sx/kpazVcPt/2BMK7DGZAsLS+tnXW8TTlFzDWau5v6yuBKFoNVzDvRHzwPubIZUgFK2GK0y4zCGMrniuVElCwWo4ozHo3vxEd+O9IoRCzslvNCSbghM9UDVRgtBwTu1O8nWPy2cWeIcUzqZzn8FQlNBwLSv9iE5gWu528MbXB+mh2Hg0P/Bf7ilCCBqUfOuEb7dwsPmBeK8gYZHcILC2h+NplYzp3nTfzd02m0cYnv6U0xd3XLxSObcPQwWR/wAgJrQjr6RU+kbRWZjwkd9ENzWE+qsh1F8Nof5qCPUXJiy9V1/Of4hbKJyXlj9voR6i/HX8V6shbAgbQvlqCBvChlC+cE6zHZbUU9/FqEdklbvOb5vUJEyoYjWwrBpC/dUQ6q+GUH81hPqrIdRfDaH+aggreEL1t3zu+a8mdMT/J6xevZzQ+1v5LZ/Tqwmd2bjqWz7bAkQYlDycJyjwN3lfv5VC2GpXrYndakn2moqriZyMbRSYnrSFxToIW93pZHXwU876+7MWwkzz3vp0sKLAqjevqJEwVbj47L9vg5izru6sm/DMud+Mjls/Mms4aCWHMJM93q0+3ODFr6dMwkzz6frPIR62ydfSfydhpkWv8z58iQ2pQpgqXGxGy+0sfj0rHLZKEWYKx7v2d3VZgoKEmSrLEpQlzFRBlqA4YapyWYIOhJnC/e6SJTzFqQ9hpuezBN0IM817k4ezBD0JMz2WJehMmCrNEoyCLEF7wkwFWcIvIczEZgm/ijBTkiUMLP/yev5CwlRplmAkWYIX3L+6pKQQZkprCe7LHyORsCY1hPqrIdRfv5/wP1HkE5RE/s+3AAAAAElFTkSuQmCC'

definition = ActionDefinition(
                name= slug,
                url= f'{action_hub}/actions/{slug}/action',
                label= 'Word',
                icon_data_uri= icon_data_uri,
                form_url= f'{action_hub}/actions/{slug}/form',
                supported_action_types= ['query'],
                description= 'This action will generate a Word document based on a Look',
                params= [],
                supported_formats= ['wysiwyg_png', 'csv'],
                supported_formattings= ['formatted', 'unformatted'],
                supported_visualization_formattings= ['apply', 'noapply'],
            )

@app.post(f'/actions/{slug}/form')
def form():
    """Standard Action Hub endpoint. Returns this action's sending/scheduling form as a JSON response."""
    return [
        ActionFormField(
            name='email_address',
            label='Email Address',
            description='Email address to send Word document',
            required=True,
        ),
        ActionFormField(
            name='email_subject',
            label='Subject',
            description='Email subject line',
            required=True,
        ),
        ActionFormField(
            name='email_body',
            label='Body',
            description='Email body text',
            required=True,
            type='textarea'
        ),
        ActionFormField(
            name='filename',
            label='Filename',
            description='Filename for the generated Word document',
            required=True,
        ),
    ]

@app.post(f'/actions/{slug}/action')
def action(payload: ActionRequest):
    """Endpoint for the Look to Word action: runs the given Look and embeds in a Word doc along with its description."""
    sdk = get_sdk_for_schedule(payload.scheduled_plan.scheduled_plan_id)
    plan = sdk.scheduled_plan(payload.scheduled_plan.scheduled_plan_id)
    look = sdk.look(plan.look_id)
    image = sdk.run_look(plan.look_id, 'png')
    description = look.description

    image_file = get_temp_file_name(
        slug, 
        '.'.join(['look', str(look.id), 'png'])
    )
    with open(image_file, 'wb') as file:
        file.write(image)
    
    document = Document()
    document.add_heading('Exported Look', 0)

    p = document.add_paragraph(description)

    document.add_picture(image_file, width=Inches(4.50))

    file_name = get_output_file_name(
        slug, 
        '.'.join(['look', str(look.id), 'docx']), 
        timestamp=True
    )
    document.save(file_name)

    response = send_email(
        to_emails=payload.form_params['email_address'],
        subject=payload.form_params['email_subject'],
        body=payload.form_params['email_body'],
        file_name=file_name,
        file_type='docx'
    )

    return {'response': 'response'}

